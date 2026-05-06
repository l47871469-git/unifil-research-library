import io
import json
import tempfile
import datetime
import streamlit as st
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))
from utils.data_utils import load_sources, save_sources, THEMATIC_CLUSTERS

SOURCE_TYPES = [
    "Academic",
    "Think tank",
    "UN Official Publication",
    "UN/NGO Statement",
    "Media",
    "Opinion",
    "Other",
]

EVENT_CATEGORIES = [
    "UNIFIL",
    "Israel-Lebanon",
    "Hezbollah",
    "Lebanese Politics",
    "Regional",
    "International",
]

SCANNER_SYSTEM = """You are a research analyst processing sources for the UNIFIL Lessons Learned exercise (UN DPO/PBPS). Extract and structure the following fields from the document provided.

Return ONLY valid JSON with exactly these keys:
{
  "id": "lastname_year (e.g. smith_2020, use unknown_year if unclear)",
  "author": "Author(s) full name(s)",
  "title": "Full title",
  "source_type": "One of: Academic / Think tank / UN Official Publication / UN/NGO Statement / Media / Opinion / Other",
  "year": integer or null,
  "publisher": "Publisher or outlet",
  "tags": ["tag1", "tag2"],
  "abstract": "100 words max. Main argument and scope.",
  "key_arguments": [
    "Argument 1",
    "Argument 2",
    "Argument 3"
  ],
  "bias_flag": "One paragraph on institutional affiliation and potential bias.",
  "thematic_clusters": [
    "Cluster name from list"
  ],
  "timeline_events": [
    {
      "date": "YYYY-MM-DD or YYYY-MM or YYYY",
      "event": "Event title",
      "category": "One of: UNIFIL / Israel-Lebanon / Hezbollah / Lebanese Politics / Regional / International",
      "description": "1-2 sentence description"
    }
  ],
  "lessons_learned": [
    {
      "text": "Lesson text",
      "tag": "SOURCE-DERIVED or ANALYTICAL INFERENCE",
      "theme": "Thematic cluster name"
    }
  ],
  "actors": ["Actor name 1", "Actor name 2"]
}

Thematic clusters to map to (use exact names):
- Mandate evolution
- Political dynamics & P5
- Liaison & tripartite mechanism
- Monitoring, reporting & technology
- TCC dynamics & command
- Host-state relations & LAF
- Civilian protection
- Force protection & safety
- Operational adaptation & innovation
- Relations with non-state armed actors
- De-mining & post-conflict stabilization
- CIMIC & community relations
- DPKO-DPPA integration

Rules:
- Never invent content not present in the source
- Timeline events: major events only (wars, massacres, mandate changes, key decisions)
- Lessons learned: 2-4 max, tag rigorously
- Return ONLY the JSON object, no preamble, no markdown fences
"""


def _init_state():
    if "add_src_stage" not in st.session_state:
        st.session_state.add_src_stage = "upload"
    if "add_src_result" not in st.session_state:
        st.session_state.add_src_result = None
    if "add_src_saved_id" not in st.session_state:
        st.session_state.add_src_saved_id = None
    if "add_src_saved_title" not in st.session_state:
        st.session_state.add_src_saved_title = None


def _scan_pdf(uploaded_file, api_key: str) -> dict:
    from pypdf import PdfReader
    import anthropic

    try:
        reader = PdfReader(io.BytesIO(uploaded_file.read()))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")

    if not text.strip():
        raise RuntimeError("Could not extract any text from the PDF. The file may be scanned/image-only.")

    try:
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            system=SCANNER_SYSTEM,
            messages=[{
                "role": "user",
                "content": f"Process this document:\n\n{text[:15000]}"
            }]
        )
        raw = response.content[0].text.strip()
        # Strip markdown fences if the model adds them despite instructions
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise RuntimeError(f"API returned invalid JSON: {e}")
    except Exception as e:
        raise RuntimeError(f"API call failed: {e}")


def _generate_docx(sources: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("UNIFIL Lessons Learned Research Corpus")
    run.bold = True
    run.font.size = Pt(20)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Generated: {datetime.date.today().strftime('%d %B %Y')}")

    doc.add_page_break()

    sorted_sources = sorted(sources, key=lambda s: s.get("year") or 0)

    for i, s in enumerate(sorted_sources):
        doc.add_heading(s.get("title", "Untitled"), level=1)

        meta = f"{s.get('author', '')}  |  {s.get('year', '')}  |  {s.get('source_type', '')}  |  {s.get('publisher', '')}"
        doc.add_paragraph(meta)

        abstract = s.get("abstract", "")
        if abstract:
            doc.add_heading("Abstract", level=2)
            doc.add_paragraph(abstract)

        key_args = s.get("key_arguments", [])
        if key_args:
            doc.add_heading("Key Arguments", level=2)
            for arg in key_args:
                doc.add_paragraph(arg, style="List Bullet")

        bias = s.get("bias_flag", "")
        if bias:
            doc.add_heading("Bias Flag", level=2)
            doc.add_paragraph(bias)

        clusters = s.get("thematic_clusters", [])
        if clusters:
            doc.add_heading("Thematic Clusters", level=2)
            doc.add_paragraph(", ".join(clusters))

        lessons = s.get("lessons_learned", [])
        if lessons:
            doc.add_heading("Lessons Learned Candidates", level=2)
            for ll in lessons:
                tag = ll.get("tag", "")
                text = ll.get("text", "")
                doc.add_paragraph(f"[{tag}] {text}", style="List Bullet")

        events = s.get("timeline_events", [])
        if events:
            doc.add_heading("Timeline Events", level=2)
            for ev in events:
                line = f"{ev.get('date', '')} — {ev.get('event', '')}: {ev.get('description', '')}"
                doc.add_paragraph(line, style="List Bullet")

        if i < len(sorted_sources) - 1:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _stage_upload():
    st.markdown("""
    <div style="font-size:0.9rem; color:#5a6a7a; margin-bottom:1.5rem; max-width:640px;">
        Upload a PDF to scan and add to the research corpus.
    </div>
    """, unsafe_allow_html=True)

    # API key input — pre-filled from env if available
    api_key = st.text_input(
        "Anthropic API key",
        value=st.session_state.add_src_api_key,
        type="password",
        placeholder="sk-ant-…",
        key="add_src_api_key_input",
        help="Your key is only held in memory for this session and never saved to disk.",
    )
    st.session_state.add_src_api_key = api_key

    uploaded = st.file_uploader(
        "Upload PDF",
        type=["pdf"],
        key="add_src_file",
        label_visibility="collapsed",
    )

    scan_ready = bool(uploaded and api_key.strip())
    col_btn, _ = st.columns([1, 4])
    with col_btn:
        scan_clicked = st.button("Scan document", key="add_src_scan_btn", disabled=not scan_ready)

    if scan_clicked and scan_ready:
        with st.spinner("Scanning document…"):
            try:
                result = _scan_pdf(uploaded, api_key.strip())
                st.session_state.add_src_result = result
                st.session_state.add_src_stage = "review"
                st.rerun()
            except RuntimeError as e:
                st.error(str(e))


def _stage_review():
    r = st.session_state.add_src_result or {}

    st.markdown("""
    <div style="font-size:0.9rem; color:#5a6a7a; margin-bottom:1.5rem; max-width:640px;">
        Review and edit the extracted fields before saving to the library.
    </div>
    """, unsafe_allow_html=True)

    # ── Row 1: id / year ──────────────────────────────────────────────────────
    col_id, col_year = st.columns(2)
    with col_id:
        new_id = st.text_input("ID", value=r.get("id", ""), key="add_src_id")
    with col_year:
        year_val = r.get("year") or 2024
        if not isinstance(year_val, int):
            try:
                year_val = int(year_val)
            except Exception:
                year_val = 2024
        new_year = st.number_input("Year", min_value=1900, max_value=2030, value=year_val, step=1, key="add_src_year")

    # ── Row 2: author / source_type ───────────────────────────────────────────
    col_author, col_type = st.columns(2)
    with col_author:
        new_author = st.text_input("Author(s)", value=r.get("author", ""), key="add_src_author")
    with col_type:
        existing_type = r.get("source_type", SOURCE_TYPES[0])
        type_idx = SOURCE_TYPES.index(existing_type) if existing_type in SOURCE_TYPES else 0
        new_type = st.selectbox("Source type", SOURCE_TYPES, index=type_idx, key="add_src_type")

    # ── Row 3: title ──────────────────────────────────────────────────────────
    new_title = st.text_input("Title", value=r.get("title", ""), key="add_src_title")

    # ── Row 4: publisher ──────────────────────────────────────────────────────
    new_publisher = st.text_input("Publisher / outlet", value=r.get("publisher", ""), key="add_src_publisher")

    # ── Row 5: abstract ───────────────────────────────────────────────────────
    new_abstract = st.text_area("Abstract", value=r.get("abstract", ""), height=120, key="add_src_abstract")

    # ── Row 6: bias_flag ──────────────────────────────────────────────────────
    new_bias = st.text_area("Bias flag", value=r.get("bias_flag", ""), height=100, key="add_src_bias")

    # ── Row 7: key arguments ─────────────────────────────────────────────────
    st.markdown('<div style="font-size:0.82rem;font-weight:600;color:#374151;margin:0.6rem 0 0.3rem 0;text-transform:uppercase;letter-spacing:0.05em;">Key Arguments</div>', unsafe_allow_html=True)
    if "add_src_args" not in st.session_state:
        st.session_state.add_src_args = list(r.get("key_arguments", []))
    args = st.session_state.add_src_args
    to_remove_arg = None
    for i, arg in enumerate(args):
        col_arg, col_rm = st.columns([9, 1])
        with col_arg:
            args[i] = st.text_area(f"Argument {i+1}", value=arg, height=68, key=f"add_src_arg_{i}", label_visibility="collapsed")
        with col_rm:
            st.markdown('<div style="margin-top:0.4rem;"></div>', unsafe_allow_html=True)
            if st.button("✕", key=f"add_src_arg_rm_{i}"):
                to_remove_arg = i
    if to_remove_arg is not None:
        st.session_state.add_src_args.pop(to_remove_arg)
        st.rerun()
    if st.button("+ Add argument", key="add_src_arg_add"):
        st.session_state.add_src_args.append("")
        st.rerun()

    # ── Row 8: thematic clusters ─────────────────────────────────────────────
    existing_clusters = [c for c in r.get("thematic_clusters", []) if c in THEMATIC_CLUSTERS]
    new_clusters = st.multiselect("Thematic clusters", THEMATIC_CLUSTERS, default=existing_clusters, key="add_src_clusters")

    # ── Row 9: tags ───────────────────────────────────────────────────────────
    new_tags_str = st.text_input("Tags (comma-separated)", value=", ".join(r.get("tags", [])), key="add_src_tags")

    # ── Row 10: timeline events ───────────────────────────────────────────────
    st.markdown('<div style="font-size:0.82rem;font-weight:600;color:#374151;margin:0.6rem 0 0.3rem 0;text-transform:uppercase;letter-spacing:0.05em;">Timeline Events</div>', unsafe_allow_html=True)
    if "add_src_events" not in st.session_state:
        st.session_state.add_src_events = list(r.get("timeline_events", []))
    events = st.session_state.add_src_events
    to_remove_ev = None
    for i, ev in enumerate(events):
        st.markdown(f'<div style="background:#fafaf8;border:1px solid #e8e5de;border-radius:4px;padding:0.6rem 0.8rem;margin-bottom:0.4rem;">', unsafe_allow_html=True)
        col_date, col_cat, col_rm = st.columns([2, 2, 0.6])
        with col_date:
            events[i]["date"] = st.text_input("Date", value=ev.get("date", ""), key=f"add_src_ev_date_{i}")
        with col_cat:
            cat_val = ev.get("category", EVENT_CATEGORIES[0])
            cat_idx = EVENT_CATEGORIES.index(cat_val) if cat_val in EVENT_CATEGORIES else 0
            events[i]["category"] = st.selectbox("Category", EVENT_CATEGORIES, index=cat_idx, key=f"add_src_ev_cat_{i}")
        with col_rm:
            st.markdown('<div style="margin-top:1.6rem;"></div>', unsafe_allow_html=True)
            if st.button("✕", key=f"add_src_ev_rm_{i}"):
                to_remove_ev = i
        events[i]["event"] = st.text_input("Event title", value=ev.get("event", ""), key=f"add_src_ev_title_{i}")
        events[i]["description"] = st.text_area("Description", value=ev.get("description", ""), height=80, key=f"add_src_ev_desc_{i}")
        st.markdown('</div>', unsafe_allow_html=True)
    if to_remove_ev is not None:
        st.session_state.add_src_events.pop(to_remove_ev)
        st.rerun()
    if st.button("+ Add event", key="add_src_ev_add"):
        st.session_state.add_src_events.append({"date": "", "event": "", "category": EVENT_CATEGORIES[0], "description": ""})
        st.rerun()

    # ── Row 11: lessons learned ───────────────────────────────────────────────
    st.markdown('<div style="font-size:0.82rem;font-weight:600;color:#374151;margin:0.6rem 0 0.3rem 0;text-transform:uppercase;letter-spacing:0.05em;">Lessons Learned Candidates</div>', unsafe_allow_html=True)
    if "add_src_lessons" not in st.session_state:
        st.session_state.add_src_lessons = list(r.get("lessons_learned", []))
    lessons = st.session_state.add_src_lessons
    to_remove_ll = None
    ll_tags = ["SOURCE-DERIVED", "ANALYTICAL INFERENCE"]
    for i, ll in enumerate(lessons):
        st.markdown('<div style="background:#fafaf8;border:1px solid #e8e5de;border-radius:4px;padding:0.6rem 0.8rem;margin-bottom:0.4rem;">', unsafe_allow_html=True)
        col_tag, col_theme, col_rm = st.columns([2, 2, 0.6])
        with col_tag:
            raw_tag = ll.get("tag", ll_tags[0])
            # Normalise tag value from scanner
            norm_tag = raw_tag.upper().replace("-", " ") if raw_tag else ll_tags[0]
            if norm_tag not in ll_tags:
                norm_tag = ll_tags[0]
            tag_idx = ll_tags.index(norm_tag)
            lessons[i]["tag"] = st.selectbox("Type", ll_tags, index=tag_idx, key=f"add_src_ll_tag_{i}")
        with col_theme:
            theme_val = ll.get("theme", THEMATIC_CLUSTERS[0])
            theme_idx = THEMATIC_CLUSTERS.index(theme_val) if theme_val in THEMATIC_CLUSTERS else 0
            lessons[i]["theme"] = st.selectbox("Theme", THEMATIC_CLUSTERS, index=theme_idx, key=f"add_src_ll_theme_{i}")
        with col_rm:
            st.markdown('<div style="margin-top:1.6rem;"></div>', unsafe_allow_html=True)
            if st.button("✕", key=f"add_src_ll_rm_{i}"):
                to_remove_ll = i
        lessons[i]["text"] = st.text_area("Lesson text", value=ll.get("text", ""), height=80, key=f"add_src_ll_text_{i}")
        st.markdown('</div>', unsafe_allow_html=True)
    if to_remove_ll is not None:
        st.session_state.add_src_lessons.pop(to_remove_ll)
        st.rerun()
    if st.button("+ Add lesson", key="add_src_ll_add"):
        st.session_state.add_src_lessons.append({"text": "", "tag": ll_tags[0], "theme": THEMATIC_CLUSTERS[0]})
        st.rerun()

    # ── Row 12: actors ────────────────────────────────────────────────────────
    new_actors_str = st.text_input("Actors (comma-separated)", value=", ".join(r.get("actors", [])), key="add_src_actors")

    st.markdown('<div style="height:1rem;"></div>', unsafe_allow_html=True)

    # ── Action buttons ────────────────────────────────────────────────────────
    col_save, col_reset, _ = st.columns([1.4, 1.2, 4])
    with col_save:
        save_clicked = st.button("💾 Save to library", key="add_src_save_btn")
    with col_reset:
        if st.button("← Start over", key="add_src_reset_btn"):
            _clear_review_state()
            st.session_state.add_src_stage = "upload"
            st.session_state.add_src_result = None
            st.rerun()

    if save_clicked:
        # Read current widget values for lists that mutate in session_state
        final_args = [st.session_state.get(f"add_src_arg_{i}", a) for i, a in enumerate(st.session_state.add_src_args)]
        final_args = [a for a in final_args if a.strip()]

        final_events = []
        for i, ev in enumerate(st.session_state.add_src_events):
            final_events.append({
                "date": st.session_state.get(f"add_src_ev_date_{i}", ev.get("date", "")),
                "event": st.session_state.get(f"add_src_ev_title_{i}", ev.get("event", "")),
                "category": st.session_state.get(f"add_src_ev_cat_{i}", ev.get("category", "")),
                "description": st.session_state.get(f"add_src_ev_desc_{i}", ev.get("description", "")),
            })

        final_lessons = []
        for i, ll in enumerate(st.session_state.add_src_lessons):
            final_lessons.append({
                "text": st.session_state.get(f"add_src_ll_text_{i}", ll.get("text", "")),
                "tag": st.session_state.get(f"add_src_ll_tag_{i}", ll.get("tag", "")),
            })

        tags_list = [t.strip() for t in new_tags_str.split(",") if t.strip()]
        actors_list = [a.strip() for a in new_actors_str.split(",") if a.strip()]

        source_id = new_id.strip() or "unknown_year"
        sources = load_sources()
        existing_ids = {s.get("id") for s in sources}
        if source_id in existing_ids:
            source_id = source_id + "_v2"

        new_source = {
            "id": source_id,
            "author": new_author.strip(),
            "title": new_title.strip(),
            "source_type": new_type,
            "year": int(new_year),
            "publisher": new_publisher.strip(),
            "tags": tags_list,
            "abstract": new_abstract.strip(),
            "key_arguments": final_args,
            "bias_flag": new_bias.strip(),
            "thematic_clusters": new_clusters,
            "timeline_events": final_events,
            "lessons_learned": final_lessons,
            "actors": actors_list,
        }

        sources.append(new_source)
        save_sources(sources)
        st.session_state.add_src_saved_id = source_id
        st.session_state.add_src_saved_title = new_title.strip()
        _clear_review_state()
        st.session_state.add_src_stage = "saved"
        st.rerun()


def _stage_saved():
    saved_title = st.session_state.get("add_src_saved_title", "")
    saved_id = st.session_state.get("add_src_saved_id", "")

    st.markdown(f"""
    <div style="background:#f0faf0;border:1px solid #9ac99a;border-left:3px solid #16a34a;
         border-radius:4px;padding:1rem 1.2rem;margin-bottom:1.2rem;">
        <div style="font-size:0.95rem;font-weight:600;color:#15803d;margin-bottom:0.3rem;">
            Source saved successfully.
        </div>
        <div style="font-size:0.88rem;color:#374151;">{saved_title}</div>
        <div style="font-size:0.78rem;color:#8a9ab0;margin-top:0.2rem;">ID: {saved_id}</div>
    </div>
    """, unsafe_allow_html=True)

    _export_button()

    st.markdown('<div style="height:0.5rem;"></div>', unsafe_allow_html=True)

    col_add, col_view, _ = st.columns([1.4, 1.4, 3])
    with col_add:
        if st.button("Add another source", key="add_src_another_btn"):
            st.session_state.add_src_stage = "upload"
            st.session_state.add_src_result = None
            st.rerun()
    with col_view:
        if st.button("View in library", key="add_src_view_btn"):
            st.session_state.page = "Sources"
            st.rerun()


def _export_button():
    sources = load_sources()
    if not sources:
        return
    try:
        docx_bytes = _generate_docx(sources)
        filename = f"UNIFIL_LL_Master_Document_{datetime.date.today().isoformat()}.docx"
        st.download_button(
            label="⬇ Export master document (.docx)",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="add_src_export_btn",
        )
    except Exception as e:
        st.error(f"Export failed: {e}")


def _clear_review_state():
    for k in ("add_src_args", "add_src_events", "add_src_lessons"):
        st.session_state.pop(k, None)


def show():
    _init_state()

    # ── Page header ───────────────────────────────────────────────────────────
    col_hdr, col_export = st.columns([4, 1.4])
    with col_hdr:
        st.markdown("""
        <div class="library-header">
            <div>
                <p class="library-subtitle">UN DPO · Policy & Best Practices Service</p>
                <h1 class="library-title">Add Source</h1>
            </div>
        </div>
        """, unsafe_allow_html=True)
    with col_export:
        st.markdown('<div style="margin-top:2.1rem;"></div>', unsafe_allow_html=True)
        _export_button()

    stage = st.session_state.add_src_stage

    if stage == "upload":
        _stage_upload()
    elif stage == "review":
        _stage_review()
    elif stage == "saved":
        _stage_saved()
